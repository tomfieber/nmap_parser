#!/usr/bin/env python3

import os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, RGBColor, Pt
from Modules.keyFunctions import join_values, split_banner


WHITE = "FFFFFF"
TABLE_STYLE = "TrustFoundry Table3"
BULLET_STYLE = "BulletedList"
SUB_BULLET_STYLE = "BulletedList 2"
TEXT_STYLE = "Text"
TESTER_NAME = "TrustFoundry"
HEADING_LEVEL = 2


def export_db(db, name, attr, doc):
    for d in db[name][attr]:
        doc.add_paragraph(d, style=SUB_BULLET_STYLE)


class AppendixGenerator:

    def __init__(self, options):
        self.options = options
        self.MAX_LIMIT = 30

    def export_doc(self, dehashed_dict, ips, ports, document):

        if not os.path.exists('./appendix/'):
            os.mkdir('./appendix/')

        if not self.options.files:
            DEHASHED_HEADING = "I"
        else:
            DEHASHED_HEADING = "III"


        if self.options.files:
            # Add the hostname mapping table
            document.add_page_break()
            document.add_heading(
                'Appendix I - Enumerated Hostnames', level=HEADING_LEVEL)
            hostname_summary = "The following table illustrates the mappings of all IP addresses and hostnames enumerated during this engagement."
            document.add_paragraph(hostname_summary, style=TEXT_STYLE)
            table = document.add_table(rows=1, cols=2)
            table.style = TABLE_STYLE
            cells = table.rows[0].cells

            run = cells[0].paragraphs[0].add_run("IP")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[1].paragraphs[0].add_run("Hostnames")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            for cell in cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)

            for ip in sorted(ips.keys(), key=lambda ip: [int(ip) for ip in ip.split('.')]):
                row = table.add_row().cells
                run = row[0].paragraphs[0].add_run(ip)
                run.font.bold = False
                for host in ips[ip]:
                    lasthost = ips[ip][-1]
                    if host is not lasthost and host != '':
                        run = row[1].paragraphs[0].add_run(host)
                    elif host != '':
                        run = row[1].paragraphs[0].add_run(host)

                    for r in row:
                        r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r.paragraphs[0].paragraph_format.space_before = Pt(2)
                        r.paragraphs[0].paragraph_format.space_after = Pt(2)

            # Add the table of listening services
            document.add_page_break()
            document.add_heading(
                'Appendix II - Listening Services', level=HEADING_LEVEL)
            summary0 = "The following table illustrates all the services TrustFoundry enumerated during this engagement."
            document.add_paragraph(summary0, style=TEXT_STYLE)
            table = document.add_table(rows=1, cols=6)
            table.style = TABLE_STYLE
            cells = table.rows[0].cells

            run = cells[0].paragraphs[0].add_run("IP")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[1].paragraphs[0].add_run("Port")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[2].paragraphs[0].add_run("Protocol")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[3].paragraphs[0].add_run("Service Type")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[4].paragraphs[0].add_run("Identified Services")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            run = cells[5].paragraphs[0].add_run("Identified Version")
            run.font.color.rgb = RGBColor.from_string(WHITE)

            for cell in cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)

            for ipaddr in sorted(ports.keys(), key=lambda ip: [int(ip) for ip in ip.split('.')]):
                for i in range(len(ports[ipaddr])):
                    row = table.add_row().cells
                    try:
                        dict = ports[ipaddr][i]
                    except KeyError:
                        continue
                    port, protocol, service = dict[
                        'port'], dict['protocol'], dict['service']
                    banner = dict['banner']
                    product, version = split_banner(banner)
                    run = row[0].paragraphs[0].add_run(ipaddr)
                    run.font.bold = False
                    run = row[1].paragraphs[0].add_run(port)
                    run = row[2].paragraphs[0].add_run(protocol)
                    run = row[3].paragraphs[0].add_run(service)
                    run = row[4].paragraphs[0].add_run(product)
                    run = row[5].paragraphs[0].add_run(version)
                    for r in row:
                        r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r.paragraphs[0].paragraph_format.space_before = Pt(
                            2)
                        r.paragraphs[0].paragraph_format.space_after = Pt(
                            2)

        if self.options.dehashed:

            """Add the table of breached credentials"""
            document.add_page_break()
            document.add_heading(
                f"Appendix {DEHASHED_HEADING} - Users with Leaked Credentials in Data Breaches", level=HEADING_LEVEL)
            summary1 = f"{TESTER_NAME} searched through various data breaches to identify any "
            summary2 = "CLIENTNAME"
            summary3 = " employees who have had credentials exposed in a breach. The following is a list of users who were identified as having had credentials exposed in one or more data breaches; note that only instances in which a user's hashed or plaintext password was identified as being exposed are included."
            para = document.add_paragraph(summary1, style=TEXT_STYLE)
            para.add_run(summary2).font.highlight_color = WD_COLOR_INDEX.YELLOW
            para.add_run(summary3)
            for name in dehashed_dict.keys():
                email = document.add_paragraph(name, style=BULLET_STYLE)
                email.paragraph_format.space_after = Inches(0)
                export_db(dehashed_dict, name, 'database', document)

        document.save('./appendix/appendix.docx')
